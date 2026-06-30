"""
resume_tailor.py — Generate a tailored .docx resume for each job application.

Usage:
    from resume_tailor import tailor_resume
    path = tailor_resume("Data Engineer", "Acme Corp", job_description, "output/resume.docx")

Requires: pip install python-docx
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Juviny's base resume content
# ---------------------------------------------------------------------------

CONTACT = {
    "name": "Your Name",
    "email": "you@example.com",
    "phone": "(555) 123-4567",
    "location": "Your City, ST",
    "linkedin": "linkedin.com/in/your-profile",
}

EDUCATION = (
    "Your Degree, Your University | Graduation Date\n"
    "Awards, honors, or relevant coursework"
)

# Skills grouped by category — reordered per role type below
SKILLS_ALL = {
    "languages": "Python, JavaScript, SQL/HQL, C++, Java",
    "frameworks": "Angular, React, FastAPI, SpaCy, PyTorch, Pandas, NumPy, SAS",
    "data": "Hadoop/Hive/Cloudera/CDP, Autosys, PowerBI, Spark, Databricks",
    "devops": "Git, Bitbucket, Docker, AWS, Jira, Salesforce",
}

# Each experience entry: (id_tag, one-line content)
# id_tag is used to select/reorder bullets per role type
EXPERIENCE = [
    ("etl",
     "Current or Recent Role — Company (Dates): "
     "Describe data, ETL, analytics, or backend accomplishments with measurable impact."),
    ("nlp",
     "ML/AI or Research Experience — Organization (Dates): "
     "Describe Python, ML, NLP, analytics, or research work relevant to target roles."),
    ("swe",
     "Software Engineering Experience — Company (Dates): "
     "Describe full-stack, backend, API, frontend, or platform engineering work."),
    ("ops",
     "Operations or Support Experience — Organization (Dates): "
     "Describe production support, quality, process improvement, or systems work."),
    ("product",
     "Product or Cross-Functional Experience — Company (Dates): "
     "Describe product delivery, stakeholder collaboration, automation, or Agile work."),
]

PROJECTS = [
    ("ruta",
     "Project One — Describe a relevant technical project and the tools used."),
    ("stock",
     "Project Two — Describe a data, backend, ML, or automation project."),
    ("plant",
     "Project Three — Describe another relevant project."),
]

LEADERSHIP = (
    "Leadership, memberships, volunteering, or certifications"
)


# ---------------------------------------------------------------------------
# Role detection → bullet ordering
# ---------------------------------------------------------------------------

def _detect_role_type(title: str, description: str) -> str:
    """Classify the role into one of our tailoring buckets."""
    combined = (title + " " + description).lower()

    if any(x in combined for x in ("data engineer", "etl", "pipeline", "hadoop", "hive", "spark", "databricks")):
        return "data_engineer"
    if any(x in combined for x in ("machine learning", "ml engineer", "ai engineer", "nlp", "natural language", "pytorch", "llm")):
        return "ml_ai"
    if any(x in combined for x in ("technical product manager", "tpm", "product manager")):
        return "tpm"
    if any(x in combined for x in ("quant", "quantitative", "financial engineer", "algo")):
        return "quant"
    if any(x in combined for x in ("robotics", "computer vision", "automation engineer", "cv engineer")):
        return "robotics_cv"
    if any(x in combined for x in ("solutions engineer", "sales engineer", "presales", "pre-sales")):
        return "solutions"
    if any(x in combined for x in ("backend", "back-end", "api", "fastapi", "rest", "microservice")):
        return "backend"
    return "default"


# Maps role_type → preferred order of EXPERIENCE id_tags
_EXPERIENCE_ORDER = {
    "data_engineer": ["etl", "nlp", "swe", "ops", "product"],
    "ml_ai":         ["nlp", "swe", "etl", "product", "ops"],
    "tpm":           ["product", "swe", "ruta_note", "etl", "ops"],
    "quant":         ["etl", "ops", "nlp", "swe", "product"],
    "robotics_cv":   ["nlp", "swe", "etl", "product", "ops"],
    "solutions":     ["swe", "product", "etl", "nlp", "ops"],
    "backend":       ["swe", "etl", "nlp", "product", "ops"],
    "default":       ["etl", "swe", "nlp", "product", "ops"],
}

# Extra skills to surface when keywords appear in JD
_EXTRA_SKILLS = {
    "spark":      "Spark",
    "databricks": "Databricks",
    "kafka":      "Kafka",
    "airflow":    "Airflow",
    "dbt":        "dbt",
    "snowflake":  "Snowflake",
    "redshift":   "Redshift",
    "kubernetes": "Kubernetes",
    "terraform":  "Terraform",
    "fastapi":    "FastAPI",
    "flask":      "Flask",
    "llm":        "LLM fine-tuning",
    "langchain":  "LangChain",
}


def _build_skills_line(role_type: str, description: str) -> str:
    """Return a single skills line, surfacing JD-matched extras first."""
    desc_lower = description.lower()
    extras = [v for k, v in _EXTRA_SKILLS.items() if k in desc_lower]

    base_skills = (
        "Python, SQL/HQL, JavaScript, C++, Java | "
        "Angular, React, FastAPI, SpaCy, PyTorch, Pandas, NumPy, SAS | "
        "Hadoop/Hive/Cloudera/CDP, Autosys, PowerBI | "
        "Git, Docker, AWS, Jira, Salesforce"
    )

    if extras:
        return f"{', '.join(extras)} | {base_skills}"
    return base_skills


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def _set_heading(doc: Document, text: str) -> None:
    """Add a bold, underlined section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.underline = True
    run.font.size = Pt(10)


def _set_body(doc: Document, text: str, bullet: bool = False) -> None:
    """Add a normal body paragraph, optionally as a bullet."""
    style = "List Bullet" if bullet else "Normal"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(9.5)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def tailor_resume(
    job_title: str,
    company: str,
    description: str,
    output_path: str,
) -> str:
    """
    Generate a tailored .docx resume and save it to output_path.

    Args:
        job_title:   Title of the target role
        company:     Hiring company name
        description: Full job description text
        output_path: Absolute path for the output .docx file

    Returns:
        The resolved output_path string.
    """
    # Ensure output directory exists
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    role_type = _detect_role_type(job_title, description)
    exp_order = _EXPERIENCE_ORDER.get(role_type, _EXPERIENCE_ORDER["default"])
    skills_line = _build_skills_line(role_type, description)

    # Build ordered experience list
    exp_map = {tag: text for tag, text in EXPERIENCE}
    ordered_exp = [exp_map[tag] for tag in exp_order if tag in exp_map]

    doc = Document()

    # --- Page margins (tighter for 1-page fit) ---
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- Header: Name ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(CONTACT["name"])
    name_run.bold = True
    name_run.font.size = Pt(16)

    # --- Contact line ---
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(2)
    contact_line = (
        f"{CONTACT['email']}  ·  {CONTACT['phone']}  ·  "
        f"{CONTACT['location']}  ·  {CONTACT['linkedin']}"
    )
    cr = contact_p.add_run(contact_line)
    cr.font.size = Pt(9)

    # --- Education ---
    _set_heading(doc, "Education")
    _set_body(doc, EDUCATION)

    # --- Skills ---
    _set_heading(doc, "Skills")
    _set_body(doc, skills_line)

    # --- Experience ---
    _set_heading(doc, "Experience")
    for entry in ordered_exp:
        _set_body(doc, entry, bullet=True)

    # --- Projects ---
    _set_heading(doc, "Projects")
    for _tag, proj_text in PROJECTS:
        _set_body(doc, proj_text, bullet=True)

    # --- Leadership ---
    _set_heading(doc, "Leadership & Activities")
    _set_body(doc, LEADERSHIP)

    doc.save(output_path)
    return output_path
