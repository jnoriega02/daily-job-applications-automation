from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).parent
OUT_DIR = BASE_DIR / "tailored_resumes"


# Replace all placeholder values with your own information before running.
CANDIDATE = {
    "name": "Your Name",
    "email": "you@example.com",
    "phone": "(555) 123-4567",
    "location": "Your City, ST",
    "linkedin": "linkedin.com/in/your-profile",
    "github": "github.com/your-username",
    "school": "Your University",
    "school_location": "City, ST",
    "degree": "Bachelor of Science in Computer Science",
    "graduation": "Graduation Date",
    "gpa": "GPA: X.X/4.0",
    "honors": "Honors, scholarships, awards, or leadership recognition",
    "coursework": "Data Structures and Algorithms, Software Engineering, Databases, Operating Systems",
}


VARIANTS = {
    "master": {
        "filename": "Your_Name_Master_Resume_LinkedIn_Indeed.docx",
        "skills": [
            ("Programming Languages", "Python, SQL, JavaScript, Java, C++"),
            ("Frameworks and Libraries", "React, Node.js, FastAPI, Flask, Pandas, NumPy"),
            ("Data, Cloud, and Tools", "AWS, Docker, Git, Jira, PowerBI"),
            ("Operating Systems", "Windows, Mac, Linux"),
        ],
        "projects": ["project_one", "project_two", "project_three"],
    },
    "software": {
        "filename": "Resume_SoftwareEngineer.docx",
        "skills": [
            ("Programming Languages", "Python, Java, JavaScript, SQL, C++"),
            ("Frameworks and Libraries", "React, Node.js, FastAPI, Flask, Pandas, NumPy"),
            ("Engineering Tools", "Git, Docker, AWS, Jira, CI/CD"),
            ("Operating Systems", "Windows, Mac, Linux"),
        ],
        "projects": ["project_one", "project_three", "project_two"],
    },
    "data": {
        "filename": "Resume_DataEngineer.docx",
        "skills": [
            ("Programming Languages", "Python, SQL, JavaScript, C++"),
            ("Data Engineering", "ETL pipelines, data validation, reporting, PowerBI, Pandas, NumPy"),
            ("Engineering Tools", "Git, Docker, AWS, Jira, Jupyter Notebook"),
            ("Operating Systems", "Windows, Mac, Linux"),
        ],
        "projects": ["project_two", "project_one", "project_three"],
    },
    "ml": {
        "filename": "Resume_MLEngineer.docx",
        "skills": [
            ("Programming Languages", "Python, SQL, JavaScript, C++"),
            ("Machine Learning and Data", "NLP, PyTorch, Pandas, NumPy, feature extraction, statistical modeling"),
            ("Frameworks and Tools", "FastAPI, React, Git, Docker, AWS, Jupyter Notebook"),
            ("Operating Systems", "Windows, Mac, Linux"),
        ],
        "projects": ["project_one", "project_two", "project_three"],
    },
}


ALIASES = {
    "Your_Name_Resume_Updated.docx": "master",
    "Resume_FinTech.docx": "data",
    "Resume_SystemsEngineer.docx": "software",
    "Resume_TPM_Solutions.docx": "master",
}


EXPERIENCE = {
    "current_role": {
        "title": "Current Software or Data Role",
        "company": "Current Company",
        "location": "City, ST",
        "date": "Month YYYY - Present",
        "software": [
            "Supported X product or workflow for Y users/business area by building Z feature, service, or automation with relevant tools.",
            "Reduced X issue or manual effort by diagnosing Y root cause and shipping Z fix through Git, Jira, and team review.",
            "Improved X release quality by adding Y validation checks and documenting Z runbook or deployment process.",
            "Delivered X workflow changes in Y environment by balancing Z reliability, defect resolution, and stakeholder communication.",
            "Improved X support handoffs by communicating Y root cause, remediation steps, and application impact to stakeholders.",
        ],
        "data": [
            "Supported X analytics workflow by designing and maintaining Y ETL/data pipeline using Z data tools.",
            "Improved X production reliability by owning Y triage process and coordinating Z fixes through Git, Jira, and Agile releases.",
            "Reduced X downstream data defects by building Y validation checks across Z ingestion, transformation, and reporting steps.",
            "Kept X reporting workflows available for Y users by documenting Z incidents, root causes, and recovery steps.",
            "Delivered X data workflow changes in Y environment by balancing Z release readiness, defect resolution, and communication.",
        ],
        "ml": [
            "Prepared X datasets for Y analytics/modeling workflows by building Z data pipelines and validation steps.",
            "Improved X data quality for Y reporting-ready workflows by validating Z inputs and resolving production defects.",
            "Reduced X manual investigation by automating Y reliability checks and release steps with Z tooling.",
            "Improved X handoffs to Y analytics partners by translating Z production issues into documentation and validation patterns.",
            "Supported X users by diagnosing Y data issues across Z ingestion, transformation, and scheduling layers.",
        ],
        "master": [
            "Supported X business workflow by building and maintaining Y software/data service with Z tools.",
            "Reduced X operational issue by diagnosing Y incidents and shipping Z fixes through Git, Jira, and team review.",
            "Improved X release quality by adding Y validation checks and documenting Z runbook or deployment process.",
            "Delivered X workflow changes in Y environment by balancing Z reliability, defect resolution, and communication.",
            "Improved X support handoffs by communicating Y root cause, remediation steps, and business impact.",
        ],
    },
    "research_or_analytics": {
        "title": "Research Assistant or Analytics Project",
        "company": "Organization",
        "location": "City, ST",
        "date": "Month YYYY - Month YYYY",
        "bullets": [
            "Produced X validated output from Y dataset by building Z Python, SQL, NLP, or analytics workflow.",
        ],
    },
    "operations": {
        "title": "Operations, Support, or Technical Assistant",
        "company": "Organization",
        "location": "City, ST",
        "date": "Month YYYY - Month YYYY",
        "bullets": [
            "Improved X data/process accuracy for Y records or users by auditing Z workflow and correcting discrepancies.",
        ],
    },
    "internship": {
        "title": "Product, Software, or Data Intern",
        "company": "Company",
        "location": "City, ST",
        "date": "Month YYYY - Month YYYY",
        "bullets": [
            "Reduced X manual work by automating Y workflow with Z tools while coordinating tasks with product and engineering partners.",
        ],
    },
}


PROJECTS = {
    "project_one": (
        "Project One",
        "React, Node.js, Firebase",
        "Improved X user workflow by building Y full-stack feature with Z technologies.",
    ),
    "project_two": (
        "Project Two",
        "Python, SQL",
        "Supported X analysis by building Y automation or CLI that processed Z data source.",
    ),
    "project_three": (
        "Project Three",
        "JavaScript, HTML, SQL",
        "Improved X data management by developing Y CRUD interface with Z database and web tools.",
    ),
}


def add_bottom_border(paragraph, color="9A9A9A", size="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_para(paragraph, before=0, after=0, line=1.0, left=0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.left_indent = Inches(left)


def add_run(paragraph, text, size=10, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def para(doc, before=0, after=0, align=None):
    p = doc.add_paragraph()
    set_para(p, before=before, after=after)
    if align is not None:
        p.alignment = align
    return p


def section(doc, title):
    p = para(doc, before=2, after=2)
    add_run(p, title, bold=True)
    add_bottom_border(p)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, left=0.24)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    add_run(p, text)


def right_tab(paragraph):
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.95), WD_TAB_ALIGNMENT.RIGHT)


def role(doc, title, company, location, date):
    p = para(doc)
    add_run(p, title, bold=True)
    add_run(p, "\t" + date)
    right_tab(p)
    p2 = para(doc)
    add_run(p2, company)
    add_run(p2, " - " + location, italic=True)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0
    return doc


def add_header(doc):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, CANDIDATE["name"], size=14, bold=True)
    p = para(doc, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(
        p,
        f"{CANDIDATE['email']} | {CANDIDATE['phone']} | {CANDIDATE['location']} | "
        f"{CANDIDATE['linkedin']} | {CANDIDATE['github']}",
    )


def add_education(doc):
    section(doc, "EDUCATION")
    p = para(doc)
    add_run(p, CANDIDATE["school"], bold=True)
    add_run(p, " - " + CANDIDATE["school_location"], italic=True)
    add_run(p, "\t" + CANDIDATE["graduation"])
    right_tab(p)
    p = para(doc)
    add_run(p, CANDIDATE["degree"], bold=True)
    add_run(p, "\t" + CANDIDATE["gpa"])
    right_tab(p)
    bullet(doc, "Honors: " + CANDIDATE["honors"])
    bullet(doc, "Relevant coursework: " + CANDIDATE["coursework"])


def add_skills(doc, skills):
    section(doc, "TECHNICAL SKILLS SUMMARY")
    for label, value in skills:
        p = para(doc)
        add_run(p, f"{label}: ", bold=True)
        add_run(p, value)


def add_experience(doc, variant):
    section(doc, "EXPERIENCE")
    current = EXPERIENCE["current_role"]
    role(doc, current["title"], current["company"], current["location"], current["date"])
    for item in current.get(variant, current["master"]):
        bullet(doc, item)

    for key in ("research_or_analytics", "operations", "internship"):
        item = EXPERIENCE[key]
        role(doc, item["title"], item["company"], item["location"], item["date"])
        for text in item["bullets"]:
            bullet(doc, text)


def add_projects(doc, project_order):
    section(doc, "PROJECTS/RESEARCH")
    for key in project_order:
        name, tech, description = PROJECTS[key]
        p = para(doc)
        add_run(p, name, bold=True)
        add_run(p, "\t" + tech)
        right_tab(p)
        bullet(doc, description)


def add_leadership(doc):
    section(doc, "LEADERSHIP/ACTIVITIES")
    p = para(doc)
    add_run(p, "Leadership Role or Technical Organization Member", bold=True)
    add_run(p, " - City, ST")
    add_run(p, "\tDates")
    right_tab(p)


def build_resume(variant="master", output_path=None):
    config = VARIANTS[variant]
    output = Path(output_path) if output_path else OUT_DIR / config["filename"]
    doc = setup_doc()
    add_header(doc)
    add_education(doc)
    add_skills(doc, config["skills"])
    add_experience(doc, variant)
    add_projects(doc, config["projects"])
    add_leadership(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def build_all():
    paths = [build_resume(name) for name in VARIANTS]
    for filename, variant in ALIASES.items():
        paths.append(build_resume(variant, OUT_DIR / filename))
    return paths


if __name__ == "__main__":
    for path in build_all():
        print(path)
